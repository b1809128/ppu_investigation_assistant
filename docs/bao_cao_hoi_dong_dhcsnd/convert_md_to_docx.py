import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_background(cell, fill_hex):
    """Set shading/background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_table_borders(table):
    """Set clean thin borders for Word table."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'CCCCCC')
            borders.append(border)
        tblPr[0].append(borders)


def format_official_document(doc, title_text):
    """Apply official Vietnamese document margins & header according to Decree 30/2020/NĐ-CP."""
    section = doc.sections[0]
    section.top_margin = Inches(0.79)     # 2.0 cm
    section.bottom_margin = Inches(0.79)  # 2.0 cm
    section.left_margin = Inches(1.18)    # 3.0 cm
    section.right_margin = Inches(0.59)   # 1.5 cm

    # Official National Header Table (2 Columns)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False

    # Col 0: Agency Header
    cell_0 = header_table.cell(0, 0)
    cell_0.width = Inches(3.2)
    p0 = cell_0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0_1 = p0.add_run("BỘ CÔNG AN\n")
    r0_1.font.name = "Times New Roman"
    r0_1.font.size = Pt(11)
    r0_1.font.bold = False
    
    r0_2 = p0.add_run("TRƯỜNG ĐẠI HỌC CẢNH SÁT NHÂN DÂN\n")
    r0_2.font.name = "Times New Roman"
    r0_2.font.size = Pt(11)
    r0_2.font.bold = True

    r0_3 = p0.add_run("------------------")
    r0_3.font.name = "Times New Roman"
    r0_3.font.size = Pt(10)

    # Col 1: National Motto
    cell_1 = header_table.cell(0, 1)
    cell_1.width = Inches(3.8)
    p1 = cell_1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1_1 = p1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    r1_1.font.name = "Times New Roman"
    r1_1.font.size = Pt(11)
    r1_1.font.bold = True

    r1_2 = p1.add_run("Độc lập - Tự do - Hạnh phúc\n")
    r1_2.font.name = "Times New Roman"
    r1_2.font.size = Pt(12)
    r1_2.font.bold = True
    r1_2.font.italic = True

    r1_3 = p1.add_run("-----------------------")
    r1_3.font.name = "Times New Roman"
    r1_3.font.size = Pt(10)

    # Spacing after header
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(12)

    # Main Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(title_text.upper())
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(15)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x1C, 0x75, 0xBB) # Navy Blue
    p_title.paragraph_format.space_after = Pt(18)


def convert_md_file_to_docx(md_path, docx_path, main_title):
    """Reads Markdown file line by line and converts to formatted Word document."""
    doc = docx.Document()
    format_official_document(doc, main_title)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_table = False
    table_lines = []

    for line in lines:
        raw_line = line.strip()

        # Skip main title if duplicate
        if raw_line.startswith("# ") and ("BÁO CÁO" in raw_line or "THUYẾT MINH" in raw_line or "HƯỚNG DẪN" in raw_line):
            continue

        # Handle Tables
        if raw_line.startswith("|") and raw_line.endswith("|"):
            in_table = True
            table_lines.append(raw_line)
            continue
        elif in_table:
            # Process buffered table lines
            if table_lines:
                rows_data = []
                for t_line in table_lines:
                    if re.match(r"^\|[\s:\-\|]+\|$", t_line):
                        continue # separator row
                    cells = [c.strip() for c in t_line.split("|")[1:-1]]
                    rows_data.append(cells)

                if rows_data:
                    num_cols = max(len(r) for r in rows_data)
                    tbl = doc.add_table(rows=len(rows_data), cols=num_cols)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    set_table_borders(tbl)

                    for r_idx, r_data in enumerate(rows_data):
                        for c_idx, cell_value in enumerate(r_data):
                            if c_idx < num_cols:
                                cell = tbl.cell(r_idx, c_idx)
                                cell.text = cell_value
                                p = cell.paragraphs[0]
                                p.paragraph_format.space_before = Pt(3)
                                p.paragraph_format.space_after = Pt(3)
                                for r in p.runs:
                                    r.font.name = "Times New Roman"
                                    r.font.size = Pt(11)

                                # Header Row Formatting
                                if r_idx == 0:
                                    set_cell_background(cell, "1C75BB")
                                    for r in p.runs:
                                        r.font.bold = True
                                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                p_sp = doc.add_paragraph()
                p_sp.paragraph_format.space_after = Pt(6)

            in_table = False
            table_lines = []

        if not raw_line:
            continue

        # Headings
        if raw_line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(raw_line[3:].strip())
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x1C, 0x75, 0xBB)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)

        elif raw_line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(raw_line[4:].strip())
            r.font.name = "Times New Roman"
            r.font.size = Pt(13)
            r.font.bold = True
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)

        elif raw_line.startswith("- ") or raw_line.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            text = raw_line[2:].strip()

            # Format bold/italic in bullet text
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = "Times New Roman"
                r.font.size = Pt(13)

            p.paragraph_format.space_after = Pt(3)

        elif raw_line.startswith("1. ") or raw_line.startswith("2. ") or raw_line.startswith("3. ") or raw_line.startswith("4. ") or raw_line.startswith("5. "):
            p = doc.add_paragraph()
            text = raw_line.strip()
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = "Times New Roman"
                r.font.size = Pt(13)

            p.paragraph_format.space_after = Pt(4)

        elif raw_line.startswith("```"):
            continue # Skip code block markers

        else:
            p = doc.add_paragraph()
            text = raw_line
            # Remove Markdown bold formatting markers
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = "Times New Roman"
                r.font.size = Pt(13)

            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.first_line_indent = Inches(0.4) # 1.0 cm

    doc.save(docx_path)
    print(f"Successfully generated: {docx_path}")


if __name__ == "__main__":
    base_dir = "/Users/quochuy/QH_Code/PPU_SYSTEM/ppu_investigation_assistant/docs/bao_cao_hoi_dong_dhcsnd"

    files_map = [
        (
            "01_dac_ta_he_thong_va_bao_cao_chuc_nang.md",
            "01_dac_ta_he_thong_va_bao_cao_chuc_nang.docx",
            "BÁO CÁO ĐẶC TẢ HỆ THỐNG VÀ BÁO CÁO CHỨC NĂNG TOÀN DIỆN"
        ),
        (
            "02_thuyet_minh_de_tai_khoa_hoc_dhcsnd.md",
            "02_thuyet_minh_de_tai_khoa_hoc_dhcsnd.docx",
            "TỜ TRÌNH THUYẾT MINH ĐỀ TÀI KHOA HỌC VÀ CÔNG NGHỆ"
        ),
        (
            "03_huong_dan_van_hanh_dieu_tra_vien_csnd.md",
            "03_huong_dan_van_hanh_dieu_tra_vien_csnd.docx",
            "HƯỚNG DẪN VẬN HÀNH NGHIỆP VỤ ĐIỀU TRA VÀ ĐỐI CHIẾU PHÁP LÝ"
        ),
        (
            "04_danh_gia_dinh_luong_legal_ai_benchmark.md",
            "04_danh_gia_dinh_luong_legal_ai_benchmark.docx",
            "BÁO CÁO KẾT QUẢ ĐÁNH GIÁ ĐỊNH LƯỢNG LEGAL AI BENCHMARK SUITE"
        )
    ]

    for md_name, docx_name, title in files_map:
        md_file = os.path.join(base_dir, md_name)
        docx_file = os.path.join(base_dir, docx_name)
        if os.path.exists(md_file):
            convert_md_file_to_docx(md_file, docx_file, title)

